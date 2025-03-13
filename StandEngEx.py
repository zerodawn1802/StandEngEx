from docx import Document
import re, os
from zipfile import ZipFile
from docx.shared import Inches

def extract_images_from_docx(docx_path, output_folder):
    with ZipFile(docx_path, 'r') as docx_file:
        for file in docx_file.namelist():
            if file.startswith('word/media/'):
                file_name = os.path.basename(file)
                if not os.path.exists(output_folder):
                    os.makedirs(output_folder)
                with open(os.path.join(output_folder, file_name), 'wb') as img_file:
                    img_file.write(docx_file.read(file))

def get_formatted_text(run):
    text = run.text
    if run.bold:
        text = f"**{text}**"
    if run.italic:
        text = f"*{text}*"
    if run.underline:
        text = f"_{text}_"
    return text

def add_formatted_text(paragraph, text, bold=False, italic=False, underline=False):
    run = paragraph.add_run(text)
    if bold:
        run.bold = True
    if italic:
        run.italic = True
    if underline:
        run.underline = True
    return run

def process_markdown_string(paragraph, markdown_text):
    # Regex to match bold, italic, and underline Markdown
    pattern = r'(\*\*[^*]+\*\*|\*[^*]+\*|_(?!_+$)[^_\s]+_)'  # Match bold (**), italic (*), and underline (_)
    
    parts = re.split(pattern, markdown_text)  # Split text at bold, italic, and underline parts
    for part in parts:
        # Check if part is bold (i.e., wrapped in **)
        if part.startswith('**') and part.endswith('**'):
            part = part[2:-2]  # Remove '**' markers
            add_formatted_text(paragraph, part, bold=True)
        # Check if part is italic (i.e., wrapped in *)
        elif part.startswith('*') and part.endswith('*'):
            part = part[1:-1]  # Remove '*' markers
            add_formatted_text(paragraph, part, italic=True)
        # Check if part is underline (i.e., wrapped in _)
        elif part.startswith('_') and part.endswith('_'):
            part = part[1:-1]  # Remove '_' markers
            add_formatted_text(paragraph, part, underline=True)
        else:
            add_formatted_text(paragraph, part)  # Add normal text

def replace_images_with_text(docx_path, output_path):
    doc = Document(docx_path)
    new_doc = Document()  # Tạo file DOCX mới
    image_count = 0  # Đếm số lượng hình ảnh
    
    for para in doc.paragraphs:
        if 'graphicData' in para._element.xml:
            image_count += 1
            new_doc.add_paragraph(f"[Hình ảnh {image_count}]")
        else:
            new_doc.add_paragraph(para.text)

    # Xử lý bảng trong tài liệu
    for table in doc.tables:
        new_table = new_doc.add_table(rows=len(table.rows), cols=len(table.columns))
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                new_cell = new_table.cell(r_idx, c_idx)
                for para in cell.paragraphs:
                    if 'graphicData' in para._element.xml:
                        image_count += 1
                        new_cell.text += f"[Hình ảnh {image_count}]"
                    else:
                        new_cell.text += para.text
    new_doc.save(output_path)

def extract_questions_with_clean_sections(doc_path):
    doc = Document(doc_path)
    questions = {}
    current_question = None
    current_answers = []
    question_types = {}
    
    # Iterate through paragraphs in the document
    for para in doc.paragraphs:
        # Match question format like "1. Which word has the underlined part..."
        question_match = re.match(r"(\d+)\. (.*)", para.text.strip())
        if question_match:
            if current_question is not None:
                # Save the last question with its answers as a single string, keeping formatting
                questions[current_question] = "\n".join(current_answers)
            
            # Start a new question and reset answers
            current_question = question_match.group(1)
            if 1 <= int(current_question) <= 22:
                question_types[current_question] = "TN"
            elif 23 <= int(current_question) <= 26:
                question_types[current_question] = "ĐS"
            elif 27 <= int(current_question) <= 28:
                question_types[current_question] = "TN"
            elif 29 <= int(current_question) <= 40:
                question_types[current_question] = "Điền"
            formatted_question = "".join([get_formatted_text(run) for run in para.runs])  # Get formatted question text
            current_answers = [formatted_question]  # Include the formatted question

        # Match answer choices (A, B, C, D)
        elif re.match(r"^[A-D]\.", para.text.strip()):
            answer = []
            for run in para.runs:
                answer.append(get_formatted_text(run))  # Get the formatted text for each part of the answer
            current_answers.append("".join(answer))  # Combine the formatted text for the answer

    # Add the last question and answers as a single string, preserving formatting
    if current_question is not None:
        questions[current_question] = "\n".join(current_answers)

    return questions, question_types

# Refined approach: Linking answers with explanations more carefully
def extract_answers_and_explanations(ans_path):
    doc = Document(ans_path)
    answers = {}
    explains = {}
    is_explanation_section = False
    current_answer = None
    current_explanation = None

    # Take answers from the table
    answers_in_table = []
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text = cell.text.strip()
                if text:
                    answers_in_table.append(text)
    # Regex tìm hướng dẫn giải: "<STT>. ..."
    answer_pattern = re.compile(r"(\d+)\.\s(.*)")  
    explain_pattern = re.compile(r"(\d+)\.\s(.*)")
    # Loop through paragraphs in the document
    for para in doc.paragraphs:
        text = para.text.strip()
        # print(text)
        # print("---------------")
        # Check if the explanation section starts
        if text.startswith("HƯỚNG DẪN GIẢI"):
            is_explanation_section = True
            continue

        # Extract answers (first 22 answers are in table)
        if not is_explanation_section:
            # Kiểm tra nếu là đáp án
            match_answer = answer_pattern.match(text)
            if match_answer:
                stt, ans = match_answer.groups()
                answers[stt] = ans  # Lưu nội dung câu trả lời
                current_answer = stt  # Lưu STT câu hiện tại
                continue  # Chuyển sang đoạn tiếp theo

            # Nếu đoạn văn tiếp theo không phải câu trả lời/hướng dẫn giải, tiếp tục nối vào nội dung hiện tại
            if current_answer and current_answer in answers:
                answers[current_answer] += f"\n{text}"  # Ghép thêm nội dung vào câu trả lời
            
        elif is_explanation_section and text:
            # Kiểm tra nếu là hướng dẫn giải
            match_explain = explain_pattern.match(text)
            if match_explain:
                stt, exp = match_explain.groups()
                exp = "".join([get_formatted_text(run) for run in para.runs])
                explains[stt] = exp  # Lưu nội dung hướng dẫn giải
                current_explanation = stt  # Cập nhật câu hiện tại
                continue  # Chuyển sang đoạn tiếp theo
            if current_explanation and current_explanation in explains:
                explains[current_explanation] += f"\n{text}"  # Ghép thêm nội dung vào hướng dẫn giải
    return answers_in_table, answers, explains

def insert_images_into_docx(input_docx, output_docx, image_folder):
    doc = Document(input_docx)
    
    # Duyệt qua tất cả các đoạn văn bản để tìm placeholder [Hình ảnh <STT>]
    for para in doc.paragraphs:
        match = re.search(r"\[Hình ảnh (\d+)\]", para.text)  # Tìm đoạn chứa [Hình ảnh <STT>]
        if match:
            image_number = match.group(1)  # Lấy số thứ tự hình ảnh
            image_path = os.path.join(image_folder, f"image{image_number}.png")
            
            if os.path.exists(image_path):  # Kiểm tra hình ảnh có tồn tại không
                run = para.add_run()
                run.add_picture(image_path, width=Inches(5))  # Chèn hình ảnh với kích thước 5 inch (có thể thay đổi)
    # Lưu tệp DOCX mới
    doc.save(output_docx)

# Chạy hàm và hiển thị kết quả
img_path = r"C:\Users\Admin\Desktop\Maru\StandEngEx\images"  # Đường dẫn file DOCX chứa hình ảnh
output_path = r"C:\Users\Admin\Desktop\Maru\StandEngEx\Đề 1 - Done.docx"  # Đường dẫn file đầu ra
doc_path = r"C:\Users\Admin\Desktop\Maru\StandEngEx\Đề 1A.docx"  # Đường dẫn file tài liệu
tmp_path = r"C:\Users\Admin\Desktop\Maru\StandEngEx\tmp.docx"
ans_path = r"C:\Users\Admin\Desktop\Maru\StandEngEx\ĐA đề 1A.docx"  # Đường dẫn file đáp án
extract_images_from_docx(doc_path, img_path)
# replace_images_with_text(doc_path, tmp_path)
questions, question_types = extract_questions_with_clean_sections(doc_path)
answers_in_table, answers, explains = extract_answers_and_explanations(ans_path)
outputs = []
for i in range(1, 41):
    if(question_types[str(i)] == "TN"):
        if(i < 23):
            outputs.append(f"{questions[str(i)]}\n")
            outputs.append("Lời giải\n")
            outputs.append(f"{ord(answers_in_table[i-1].split(' ')[1]) - ord('A') + 1}\n")
            outputs.append("####\n")
            if(i < 10):
                outputs.append(f"{explains[str(i)][7:]}\n\n")
            else:
                outputs.append(f"{explains[str(i)][8:]}\n\n")
            
        else:
            outputs.append(f"{questions[str(i)]}\n")
            outputs.append("Lời giải\n")
            outputs.append(f"{ord(answers[str(i)]) - ord('A') + 1}\n")
            outputs.append("####\n")
            outputs.append(f"{explains[str(i)][8:]}\n\n")
    elif(question_types[str(i)] == "ĐS"):
        outputs.append(f"{questions[str(i)]}\n")
        outputs.append("Lời giải\n")
        if(answers[str(i)] == "True"):
            outputs.append("1\n")
        elif(answers[str(i)] == "False"):
            outputs.append("0\n")
        outputs.append("####\n")
        outputs.append(f"{explains[str(i)][8:]}\n\n")
    elif(question_types[str(i)] == "Điền"):
        outputs.append(f"{questions[str(i)]}\n")
        a = []
        if(i < 37):
            a = answers[str(i)].split("/")
        else:
            a = answers[str(i)].split("ý: ")[1].split("/")
        tmp = "Đáp án: ["
        for j in range(len(a)):
            if(j != len(a)-1):
                tmp += f"[{a[j]}]|"
            else:
                tmp += f"[{a[j]}]]"
        outputs.append(f"{tmp}\n")
        outputs.append("Lời giải\n")
        outputs.append(f"{explains[str(i)][8:]}\n\n")
new_doc = Document()
paragraph = new_doc.add_paragraph()
tmpdoc = "".join(str(x) for x in outputs)
process_markdown_string(paragraph, tmpdoc)
new_doc.save(output_path)
